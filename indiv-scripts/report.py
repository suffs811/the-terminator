#!/usr/bin/python3
# author: suffs811
# https://github.com/suffs811/the-terminator.git
# purpose: script for creating report file in .txt and .docx formats from terminator findings
#
# usage: python3 report.py -o <output_file_name>


import os
import sys
import argparse


parser = argparse.ArgumentParser(description="script for making a report from pentest scripts\npython3 report.py -o <output_file_name>")
parser.add_argument("-o", "--output", help="specify name for report")
args = parser.parse_args()
output = args.output


# add contents from enum.txt, priv.txt, and data_exfil.txt to file with -o output file name
def report(output):
   # get target ip from enum.txt
   ipf = open("/terminator/enum.txt")
   ipread = ipf.readline()
   ipsplit = ipf.split(" ")
   ip = ipsplit[-2].strip()

   os.system("touch /terminator/report.txt")
   os.system("echo '-+- Penetration Testing Report for {} -+-' > /terminator/report.txt".format(ip))
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '+ + + Enumeration + + +' >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("cat /terminator/enum.txt >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '+ + + Exploitation / Initial Shell + + +' >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system(" echo '*** ADD YOUR EXPLOITION METHOD FOR THE INITAL SHELL HERE ***' >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '+ + + Privilege Escalation + + +' >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("cat /terminator/priv.txt >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '+ + + Persistence and Data Exfiltration + + +' >> /terminator/report.txt")
   os.system("cat /terminator/data_exfil.txt >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '' >> /terminator/report.txt")
   os.system("echo '--- END OF REPORT ---' >> /terminator/report.txt")
   os.system("mv /terminator/report.txt /terminator/{}".format(output))
   print("### penetration test report for {} is ready at /terminator/{} ###\nplease add your method for gaining the initial shell in the '+ + + Stage 2 - Exploitation / Initial Shell + + +' section.\nall reference data for enumeration, privilege escalation, and persistence/data exfiltration are located in /terminator/ as enum.txt, priv.txt, and data_exfil.txt, respectively.".format(ip,output))
   print("\n\n-+- {} has been terminated -+-".format(ip))
   ipf.close()


# check if docx is installed on machine
def lib_check():
   try:
      import docx
      "docx" in sys.modules
   except:
      return False


# make Word (docx) file and fill with contents from terminator.py output
def doc_make(output):
   # import docx library
   from docx import Document

   # get target ip from enum.txt
   ipf = open("/terminator/enum.txt")
   ipread = ipf.readline()
   ipsplit = ipf.split(" ")
   ip = ipsplit[-2].strip()

   # get output name
   rsplit = output.split("/")
   fname = rsplit[-1]
   fsplit = fname.split(".")
   cut = fsplit[0].strip()

   txt = re.search("txt\Z", output)
   dot = re.findall("[.]", output)
   if dot.len() > 1:
      if txt:
         cut = output[:-4]
      else:
         cut = output
   elif dot.len() == 1:
      fsplit = output.split(".")
      cut = fsplit[0].strip()
   else:
      cut = output

   # create and fill document
   document = Document()
   e = open("/terminator/enum.txt")
   p = open("/terminator/priv.txt")
   x = open("/terminator/data_exfil.txt")
   ee = e.read()
   pp = p.read()
   xx = x.read()

   document.add_heading("Penetration Testing Report for {}".format(ip), 0)
   document.add_heading("Enumeration", level=1)
   document.add_paragraph(ee)
   document.add_page_break()
   document.add_heading("Exploitation / Initial Shell", level=1)
   document.add_paragraph("*** ADD YOUR EXPLOITION METHOD FOR THE INITAL SHELL HERE ***")
   document.add_page_break()
   document.add_heading("Privilege Escalation", level=1)
   document.add_paragraph(pp)
   document.add_page_break()
   document.add_heading("Persistence and Data Exfiltration", level=1)
   document.add_paragraph(xx)
   document.save("{}.docx".format(cut))
   os.system("mv {}.docx /terminator/{}.docx".format(cut,cut))

   e.close()
   p.close()
   x.close()
   ipf.close()

   print("-+- Word document saved to /terminator/{}.docx -+-".format(cut))


# call report functions
report(output)
lib = lib_check()
if lib:
   doc_make(output)
else:
   print("\n*** 'python-docx' is not installed on your machine; please run 'pip install python-docx' in your terminal and rerun report.py ***")
